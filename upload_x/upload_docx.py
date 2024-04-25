import docx
from tkinter import filedialog

def upload_docx():
    file_path = filedialog.askopenfilename(filetypes=[("DOCX Files", "*.docx")])
    if file_path:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    return None
